from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


source = Path(sys.argv[1])
output = Path(sys.argv[2])
diagram = Path(sys.argv[3])
document = Document(source)


def find_paragraph(text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def insert_before(target, text, style="Normal"):
    return target.insert_paragraph_before(text, style=style)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def create_decimal_numbering():
    numbering = document.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in [("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "left")]:
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        level.append(element)
    ppr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(num_id_element)
    ppr.append(numpr)


def paragraphs_between(start, end):
    paragraphs = document.paragraphs
    start_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._element is start._element)
    end_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._element is end._element)
    return paragraphs[start_index + 1:end_index]


architecture_heading = find_paragraph("System Architecture")
caption = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.style and paragraph.style.name == "Caption" and "Diagram architecture" in paragraph.text
)

# The diagram is the drawing-only paragraph immediately before its caption.
caption_index = next(i for i, paragraph in enumerate(document.paragraphs) if paragraph._element is caption._element)
image_paragraph = document.paragraphs[caption_index - 1]
blips = image_paragraph._element.xpath(".//a:blip")
if not blips:
    raise ValueError("Figure 8 image was not found before its caption")
relationship_id = blips[0].get(qn("r:embed"))
image_part = image_paragraph.part.related_parts[relationship_id]
image_part._blob = diagram.read_bytes()

# Resize Figure 8 to fit the A4 text area while retaining the 16:10 aspect ratio.
cx = int(6.15 * 914400)
cy = int(cx / 1.6)
for extent in image_paragraph._element.xpath(".//wp:extent"):
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))
for extent in image_paragraph._element.xpath(".//a:xfrm/a:ext"):
    extent.set("cx", str(cx))
    extent.set("cy", str(cy))

# Replace the architecture prose but preserve the existing heading, figure, and caption.
for paragraph in paragraphs_between(architecture_heading, image_paragraph):
    remove_paragraph(paragraph)

architecture_content = [
    ("Normal", "LOTUSMILE adopts a four-tier architecture that separates the user interface, request gateway, application services, and data/retrieval resources. Internal services communicate through REST APIs or database connections, while Cohere and the payment gateways remain external services accessed securely over HTTPS. Docker Compose provides isolated service deployment and persistent storage for MySQL and Qdrant."),
    ("Normal", "Tier 1 — Presentation Layer (Client): The customer-facing interface is implemented with Blade, HTML, CSS, JavaScript, Bootstrap, and jQuery. The floating chatbot is available on the home page and internal client pages. Browser requests, including POST /api/chatbot/message, are sent through the HTTP entry point rather than directly to internal services."),
    ("Normal", "Tier 2 — Gateway Layer (Nginx): Nginx acts as the reverse proxy for the platform. It serves static assets and forwards application requests to Laravel, providing a single boundary between public clients and internal backend services."),
    ("Normal", "Tier 3 — Application and AI Orchestration Layer: Laravel is the primary application service and coordinates authentication, tour management, bookings, payments, recommendations, and chatbot requests. Two AI paths operate in this layer:"),
    ("List Paragraph", "Laravel Application: ChatbotController validates incoming messages, while HybridTourRetriever coordinates semantic retrieval, SQL keyword retrieval, Reciprocal Rank Fusion (RRF), authoritative record loading, and prompt construction. The same application handles payment redirects and IPN callbacks through the Ngrok development tunnel."),
    ("List Paragraph", "Flask Recommendation Service: The existing recommendation API applies TF-IDF vectorization and cosine similarity to return tours related to a selected tour. This item-to-item recommendation path is independent of the conversational RAG chatbot."),
    ("Normal", "Tier 4 — Data and Retrieval Layer: MySQL is the authoritative transactional database for users, tours, prices, availability, bookings, payments, reviews, images, and itineraries. Qdrant stores a derived vector index of published tour content for semantic nearest-neighbor search. Cohere is an external AI provider used to create multilingual query/document embeddings and generate grounded Vietnamese responses. If Qdrant is unavailable, Laravel continues with MySQL keyword retrieval."),
]
for style, text in architecture_content:
    insert_before(image_paragraph, text, style)

# Use a fixed caption because this report already materializes the figure number as 8.
caption.text = "Figure 8. LOTUSMILE architecture with hybrid RAG subsystem"
caption.style = "Caption"

# Replace the original RAG subsection with a concise, step-oriented explanation.
rag_heading = find_paragraph("Hybrid RAG Chatbot Subsystem")
next_heading = find_paragraph("Use Case Analysis and Design")
for paragraph in [rag_heading] + paragraphs_between(rag_heading, next_heading):
    remove_paragraph(paragraph)

rag_content = [
    ("Heading 4", "Hybrid RAG Chatbot Subsystem"),
    ("Normal", "The hybrid RAG subsystem enables users to describe travel needs in natural Vietnamese and receive answers grounded in the current LOTUSMILE catalog. It complements the Flask recommendation engine: Flask recommends tours similar to a selected item, whereas RAG retrieves tours relevant to a free-form conversational query."),
    ("Normal", "Indexing flow: The Artisan command rag:sync-tours --recreate reads each published tour and its itinerary from MySQL, removes HTML formatting, creates a semantic document, requests a search_document embedding from Cohere, and stores the resulting vector and tour identifier in the persistent Qdrant collection lotusmile_tours. The current index contains 24 published tours."),
    ("Normal", "Runtime retrieval and response flow:"),
    ("RAG Numbered", "The client sends the user's message to POST /api/chatbot/message."),
    ("RAG Numbered", "ChatbotController validates the request and passes the message to HybridTourRetriever."),
    ("RAG Numbered", "Cohere converts the message into a search_query embedding, and Qdrant returns the nearest tour identifiers by cosine similarity."),
    ("RAG Numbered", "MySQL independently searches exact terms across title, destination, description, and tour category. RRF merges the semantic and lexical rankings without requiring their raw scores to share the same scale."),
    ("RAG Numbered", "Laravel reloads up to four published tour records from MySQL so that names, prices, duration, availability, destinations, and detail links come from the authoritative database rather than the vector payload."),
    ("RAG Numbered", "The retrieved context and grounding rules are sent to the Cohere chat model, which returns a concise Vietnamese response and structured tour cards for the client interface."),
    ("Normal", "Grounding and resilience: The prompt explicitly prohibits invented tours, prices, or schedules. When no relevant tour is found, the assistant asks the user to clarify the request. Embedding or Qdrant failures are logged and automatically fall back to SQL keyword retrieval, allowing the consultation feature to remain available with reduced semantic capability."),
]
rag_num_id = create_decimal_numbering()
for style, text in rag_content:
    if style == "RAG Numbered":
        paragraph = insert_before(next_heading, text, "List Paragraph")
        apply_numbering(paragraph, rag_num_id)
    else:
        insert_before(next_heading, text, style)

# Refresh TOC and figure fields when Word opens the document.
settings = document.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

document.save(output)
