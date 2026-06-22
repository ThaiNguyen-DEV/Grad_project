from hashlib import sha256
from pathlib import Path
import sys

from docx import Document
from docx.oxml.ns import qn


path = Path(sys.argv[1])
diagram = Path(sys.argv[2])
document = Document(path)
paragraphs = document.paragraphs
start = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "System Architecture")
end = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "Use Case Analysis and Design")
for i in range(start, end + 1):
    paragraph = paragraphs[i]
    text = " ".join(paragraph.text.split())
    drawing = bool(paragraph._element.xpath(".//a:blip"))
    print(f"{i}: [{paragraph.style.name}] {'<FIGURE>' if drawing else text}")
    if drawing:
        blip = paragraph._element.xpath(".//a:blip")[0]
        image_part = paragraph.part.related_parts[blip.get(qn("r:embed"))]
        print("FIGURE_MATCH=", sha256(image_part.blob).hexdigest() == sha256(diagram.read_bytes()).hexdigest())

numbered = [p for p in paragraphs[start:end] if p._p.xpath("./w:pPr/w:numPr")]
print("NUMBERED_STEPS=", len(numbered))
print("TABLES=", len(document.tables), "SECTIONS=", len(document.sections))
